"""
Tests for the de-anonymization feature (deanonymize-document change).

Covers:
- _write_reversal_sidecar: JSON structure, deduplication, opt-out
- load_reversal_sidecar: valid load, missing file, collision warning
- deanonymize_docx: round-trip via anonymize_docx + deanonymize_docx
- deanonymize_pdf: round-trip via anonymize_pdf + deanonymize_pdf
- CLI deanonymize subcommand: success, missing file, missing sidecar, bad extension
"""
import pytest
pytest.skip("Obsolete: Deanonymize now uses Excel instead of JSON sidecars", allow_module_level=True)

import json
import logging
from pathlib import Path
from unittest.mock import patch, MagicMock

from typer.testing import CliRunner

from anonymizer.replacer import (
    anonymize_docx,
    deanonymize_docx,
)
from anonymizer.cli import app


# ── Fixtures ─────────────────────────────────────────────────────────────────

SAMPLE_MAPPING = {
    "Juan Pérez": "[PERSONA_1]",
    "Acme Corp": "[ORG_1]",
}
SAMPLE_MODES = {
    "Juan Pérez": "palabra",
    "Acme Corp": "substring",
}


# ── _write_reversal_sidecar ──────────────────────────────────────────────────

class TestWriteReversalSidecar:
    def test_correct_json_structure(self, tmp_path):
        """Sidecar contains expected schema fields."""
        output = tmp_path / "out.docx"
        _write_reversal_sidecar(output, "source.docx", SAMPLE_MAPPING, SAMPLE_MODES)

        sidecar = tmp_path / "out.docx.reversal.json"
        assert sidecar.exists()

        data = json.loads(sidecar.read_text(encoding="utf-8"))
        assert data["schema_version"] == 1
        assert data["source_document"] == "source.docx"
        assert isinstance(data["replacements"], list)
        assert len(data["replacements"]) == 2

        entry_keys = {e["original"] for e in data["replacements"]}
        assert "Juan Pérez" in entry_keys
        assert "Acme Corp" in entry_keys

    def test_deduplication(self, tmp_path):
        """Duplicate (original, pseudonym) pairs are written only once."""
        dup_mapping = {
            "Juan Pérez": "[PERSONA_1]",
            "Juan Pérez": "[PERSONA_1]",  # noqa: F601 – intentional duplicate for test
        }
        output = tmp_path / "out.docx"
        _write_reversal_sidecar(output, "source.docx", dup_mapping, None)

        data = json.loads((tmp_path / "out.docx.reversal.json").read_text(encoding="utf-8"))
        assert len(data["replacements"]) == 1

    def test_opt_out_suppresses_file(self, tmp_path):
        """write_reversal=False must not create a sidecar file."""
        from anonymizer.replacer import anonymize_docx

        # We test via anonymize_docx with write_reversal=False.
        # Patch the actual DOCX writing so we don't need a real file.
        with patch("anonymizer.replacer.Path.mkdir"), \
             patch("docx.Document") as mock_doc_cls:
            mock_doc = MagicMock()
            mock_doc.sections = []
            mock_doc_cls.return_value = mock_doc

            output = tmp_path / "out.docx"
            anonymize_docx("fake.docx", output, SAMPLE_MAPPING, SAMPLE_MODES, write_reversal=False)

            sidecar = Path(str(output) + ".reversal.json")
            assert not sidecar.exists()

    def test_modes_recorded_in_sidecar(self, tmp_path):
        """Each replacement entry captures the match_mode used."""
        output = tmp_path / "out.docx"
        _write_reversal_sidecar(output, "src.docx", SAMPLE_MAPPING, SAMPLE_MODES)

        data = json.loads((tmp_path / "out.docx.reversal.json").read_text(encoding="utf-8"))
        modes_in_sidecar = {e["original"]: e["match_mode"] for e in data["replacements"]}
        assert modes_in_sidecar["Juan Pérez"] == "palabra"
        assert modes_in_sidecar["Acme Corp"] == "substring"


# ── load_reversal_sidecar ────────────────────────────────────────────────────

class TestLoadReversalSidecar:
    def _make_sidecar(self, path: Path, replacements: list[dict]):
        data = {
            "schema_version": 1,
            "source_document": "src.docx",
            "replacements": replacements,
        }
        path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

    def test_valid_sidecar_returns_correct_dict(self, tmp_path):
        """load_reversal_sidecar returns {pseudonym: original} from valid file."""
        sidecar = tmp_path / "out.docx.reversal.json"
        self._make_sidecar(sidecar, [
            {"original": "Juan Pérez", "pseudonym": "[PERSONA_1]", "match_mode": "palabra"},
            {"original": "Acme Corp",  "pseudonym": "[ORG_1]",     "match_mode": "substring"},
        ])

        result = load_reversal_sidecar(sidecar)
        assert result == {"[PERSONA_1]": "Juan Pérez", "[ORG_1]": "Acme Corp"}

    def test_missing_file_raises_file_not_found(self, tmp_path):
        """load_reversal_sidecar raises FileNotFoundError if file absent."""
        with pytest.raises(FileNotFoundError, match="not found"):
            load_reversal_sidecar(tmp_path / "nonexistent.reversal.json")

    def test_collision_logs_warning_and_uses_last_value(self, tmp_path, caplog):
        """Pseudonym collision logs WARNING and keeps last entry."""
        sidecar = tmp_path / "out.docx.reversal.json"
        self._make_sidecar(sidecar, [
            {"original": "Juan Pérez", "pseudonym": "[PERSONA_1]", "match_mode": "palabra"},
            {"original": "Juan García", "pseudonym": "[PERSONA_1]", "match_mode": "palabra"},
        ])

        with caplog.at_level(logging.WARNING, logger="anonymizer.replacer"):
            result = load_reversal_sidecar(sidecar)

        assert result["[PERSONA_1]"] == "Juan García"  # last-entry-wins
        assert any("collision" in rec.message.lower() for rec in caplog.records)

    def test_empty_replacements_returns_empty_dict(self, tmp_path):
        """Empty replacements list returns empty dict without error."""
        sidecar = tmp_path / "out.docx.reversal.json"
        self._make_sidecar(sidecar, [])
        assert load_reversal_sidecar(sidecar) == {}


# ── DOCX round-trip ──────────────────────────────────────────────────────────

class TestDeanonymizeDocxRoundTrip:
    """Integration test: anonymize then deanonymize a real DOCX and assert text equality."""

    def test_round_trip_restores_original_text(self, tmp_path):
        """anonymize_docx + deanonymize_docx round-trip restores original paragraph text."""
        from docx import Document

        # 1. Create a minimal DOCX with known content
        original_text = "Juan Pérez trabaja en Acme Corp desde el martes."
        src = tmp_path / "original.docx"
        doc = Document()
        doc.add_paragraph(original_text)
        doc.save(str(src))

        # 2. Anonymize
        anon_out = tmp_path / "anon.docx"
        mapping = {"Juan Pérez": "[PERSONA_1]", "Acme Corp": "[ORG_1]"}
        modes = {"Juan Pérez": "palabra", "Acme Corp": "substring"}
        anonymize_docx(src, anon_out, mapping, modes, write_reversal=True)

        # 3. Verify sidecar was created
        sidecar = Path(str(anon_out) + ".reversal.json")
        assert sidecar.exists(), "Sidecar must be written alongside anonimizado output"

        # 4. Verify anonymized text differs
        anon_doc = Document(str(anon_out))
        anon_text = " ".join(p.text for p in anon_doc.paragraphs)
        assert "[PERSONA_1]" in anon_text
        assert "Juan Pérez" not in anon_text

        # 5. Deanonymize
        restored_out = tmp_path / "restored.docx"
        deanonymize_docx(anon_out, restored_out, sidecar)

        # 6. Assert original text is back
        restored_doc = Document(str(restored_out))
        restored_text = " ".join(p.text for p in restored_doc.paragraphs)
        assert "Juan Pérez" in restored_text
        assert "Acme Corp" in restored_text
        assert "[PERSONA_1]" not in restored_text
        assert "[ORG_1]" not in restored_text

    def test_empty_sidecar_produces_unchanged_text(self, tmp_path):
        """deanonymize_docx with empty sidecar leaves text unchanged."""
        from docx import Document

        src = tmp_path / "anon.docx"
        doc = Document()
        doc.add_paragraph("[PERSONA_1] trabaja en [ORG_1].")
        doc.save(str(src))

        # Empty sidecar
        sidecar = tmp_path / "anon.docx.reversal.json"
        sidecar.write_text(json.dumps({"schema_version": 1, "source_document": "x", "replacements": []}), encoding="utf-8")

        out = tmp_path / "restored.docx"
        deanonymize_docx(src, out, sidecar)

        result_doc = Document(str(out))
        result_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "[PERSONA_1]" in result_text  # unchanged


# ── PDF round-trip ───────────────────────────────────────────────────────────

class TestDeanonymizePdfRoundTrip:
    """Integration test: anonymize then deanonymize a PDF and assert text content."""

    def test_round_trip_restores_original_text(self, tmp_path):
        """anonymize_pdf + deanonymize_pdf round-trip restores original text."""
        import pdfplumber
        from anonymizer.replacer import anonymize_pdf, deanonymize_pdf
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet

        original_text = "Contactar a María López de Servicios Global."

        # 1. Create source PDF via reportlab
        src = tmp_path / "original.pdf"
        doc = SimpleDocTemplate(str(src), pagesize=A4)
        styles = getSampleStyleSheet()
        doc.build([Paragraph(original_text, styles["Normal"])])

        # 2. Anonymize
        mapping = {"María López": "[PERSONA_1]", "Servicios Global": "[ORG_1]"}
        modes = {"María López": "palabra", "Servicios Global": "substring"}
        anon_out = tmp_path / "anon.pdf"
        anonymize_pdf(src, anon_out, mapping, modes, write_reversal=True)

        sidecar = Path(str(anon_out) + ".reversal.json")
        assert sidecar.exists()

        # 3. Verify anonymized content
        with pdfplumber.open(str(anon_out)) as pdf:
            anon_content = " ".join(page.extract_text() or "" for page in pdf.pages)
        assert "[PERSONA_1]" in anon_content

        # 4. Deanonymize
        restored_out = tmp_path / "restored.pdf"
        deanonymize_pdf(anon_out, restored_out, sidecar)

        # 5. Check restored text
        with pdfplumber.open(str(restored_out)) as pdf:
            restored_content = " ".join(page.extract_text() or "" for page in pdf.pages)
        assert "María López" in restored_content
        assert "Servicios Global" in restored_content
        assert "[PERSONA_1]" not in restored_content


# ── CLI integration tests ────────────────────────────────────────────────────

class TestCliDeanonymize:
    runner = CliRunner()

    def test_cli_deanonymize_docx_success(self, tmp_path):
        """CLI deanonymize uses auto-detected sidecar and exits 0."""
        from docx import Document

        src = tmp_path / "anon.docx"
        doc = Document()
        doc.add_paragraph("[PERSONA_1] firmó el contrato.")
        doc.save(str(src))

        sidecar = Path(str(src) + ".reversal.json")
        sidecar.write_text(json.dumps({
            "schema_version": 1,
            "source_document": "original.docx",
            "replacements": [{"original": "Pedro Rodríguez", "pseudonym": "[PERSONA_1]", "match_mode": "substring"}],
        }), encoding="utf-8")

        out = tmp_path / "restored.docx"
        result = self.runner.invoke(app, ["deanonymize", str(src), str(out)])

        assert result.exit_code == 0, result.output
        assert out.exists()

    def test_cli_deanonymize_missing_input(self, tmp_path):
        """CLI exits non-zero when input file does not exist."""
        result = self.runner.invoke(app, ["deanonymize", str(tmp_path / "nope.docx"), str(tmp_path / "out.docx")])
        assert result.exit_code != 0
        assert "no encontrado" in result.output.lower()

    def test_cli_deanonymize_missing_sidecar(self, tmp_path):
        """CLI exits non-zero when sidecar is absent and --reversal not provided."""
        src = tmp_path / "anon.docx"
        from docx import Document
        Document().save(str(src))

        out = tmp_path / "out.docx"
        result = self.runner.invoke(app, ["deanonymize", str(src), str(out)])
        assert result.exit_code != 0
        assert "reversi" in result.output.lower()

    def test_cli_deanonymize_unsupported_extension(self, tmp_path):
        """CLI exits non-zero for unsupported file extension."""
        src = tmp_path / "doc.txt"
        src.write_text("contenido")

        result = self.runner.invoke(app, ["deanonymize", str(src), str(tmp_path / "out.txt")])
        assert result.exit_code != 0
        assert "no soportado" in result.output.lower()

    def test_cli_deanonymize_explicit_reversal_path(self, tmp_path):
        """CLI accepts --reversal pointing to a sidecar in non-default location."""
        from docx import Document

        src = tmp_path / "anon.docx"
        doc = Document()
        doc.add_paragraph("[ORG_1] firmó.")
        doc.save(str(src))

        # Sidecar in non-default path
        sidecar = tmp_path / "custom.reversal.json"
        sidecar.write_text(json.dumps({
            "schema_version": 1,
            "source_document": "original.docx",
            "replacements": [{"original": "Empresa SA", "pseudonym": "[ORG_1]", "match_mode": "substring"}],
        }), encoding="utf-8")

        out = tmp_path / "restored.docx"
        result = self.runner.invoke(app, ["deanonymize", str(src), str(out), "--reversal", str(sidecar)])
        assert result.exit_code == 0, result.output
        assert out.exists()
