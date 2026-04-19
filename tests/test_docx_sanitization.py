import os
from pathlib import Path
from docx import Document
from docx.oxml.shared import OxmlElement, qn

from anonymizer.replacer import anonymize_docx

def test_docx_deep_sanitization(tmp_path: Path):
    doc_path = tmp_path / "dirty.docx"
    clean_path = tmp_path / "clean.docx"

    # Create a dirty document manually using lxml within docx
    doc = Document()

    # 1. Metadata
    doc.core_properties.author = "Confidential Author"
    doc.core_properties.last_modified_by = "Confidential Author"
    doc.core_properties.comments = "Secret comment"

    # 2. Add paragraph with Hyperlink, w:del, and w:ins
    p = doc.add_paragraph("Start. ")
    
    # Add a Hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    run_hl = OxmlElement('w:r')
    t_hl = OxmlElement('w:t')
    t_hl.text = "secret@email.com"
    run_hl.append(t_hl)
    hyperlink.append(run_hl)
    p._p.append(hyperlink)

    # Add Deleted text (Track Changes)
    del_node = OxmlElement('w:del')
    del_run = OxmlElement('w:r')
    del_t = OxmlElement('w:delText')
    del_t.text = "Delete"
    del_run.append(del_t)
    del_node.append(del_run)
    p._p.append(del_node)

    # Add Inserted text (Track Changes)
    ins_node = OxmlElement('w:ins')
    ins_run = OxmlElement('w:r')
    ins_t = OxmlElement('w:t')
    ins_t.text = "Insert"
    ins_run.append(ins_t)
    ins_node.append(ins_run)
    p._p.append(ins_node)

    doc.save(str(doc_path))

    # Anonymize it
    mapping = {"secret@email.com": "[CORREO_1]"}
    anonymize_docx(doc_path, clean_path, mapping)

    # Verify the cleaned doc
    clean_doc = Document(str(clean_path))
    
    # Assert metadata purged
    assert clean_doc.core_properties.author == "Anónimo"
    assert clean_doc.core_properties.last_modified_by == "Anónimo"
    assert clean_doc.core_properties.comments == ""

    # Assert content
    xml = clean_doc._element.xml
    assert "secret@email.com" not in xml
    assert "Confidential Author" not in xml
    assert "<w:hyperlink" not in xml
    assert "<w:del" not in xml
    assert "<w:ins" not in xml
    
    # Check that "Delete" was removed completely
    assert "Delete" not in xml
    
    # Check that "Insert" was kept (as normal text without w:ins)
    assert "Insert" in xml
    
    # Check that mapping was applied
    assert "[CORREO_1]" in xml

