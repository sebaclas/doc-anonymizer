"""
Applies the pseudonym mapping to DOCX and PDF documents,
generating new anonymized files.

De-anonymization (restoration) is performed using the Excel mapping file
generated during the anonymization session.
"""
import json
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

logger = logging.getLogger(__name__)


def _apply_mapping_to_text(
    text: str,
    mapping: dict[str, str],
    modes: dict[str, str] | None = None,
) -> str:
    """
    Replace all occurrences of original->pseudo in text (longest match first).

    modes: optional dict {original: match_mode} where match_mode is:
      "palabra"   - word-boundary match (\\b on both sides). Prevents partial
                    matches like "Ana" inside "Mariana". Works correctly even
                    with surrounding punctuation like parentheses or hyphens.
      "substring" - plain substring match anywhere in the text, regardless of
                    surrounding characters. Use for acronyms, codes, or names
                    that may appear embedded inside other words.
    When modes is None, defaults to pure substring for all keys (legacy behaviour).
    """
    if not mapping:
        return text

    # Sort by length descending to avoid partial replacements (longer keys win)
    sorted_keys = sorted(mapping.keys(), key=len, reverse=True)

    # Build per-key pattern parts
    parts = []
    for k in sorted_keys:
        escaped = re.escape(k)
        mode = modes.get(k, "palabra") if modes is not None else "substring"
        if mode == "palabra":
            parts.append(r"\b" + escaped + r"\b")
        else:
            parts.append(escaped)

    # Compilar con re.IGNORECASE para que no importe si es 'Grossi' o 'grossi'
    pattern = re.compile("|".join(parts), re.UNICODE | re.IGNORECASE)
    return pattern.sub(lambda m: mapping.get(m.group(0), mapping.get(m.group(0).lower(), mapping.get(m.group(0).title(), next(iter(mapping.values()))))), text)


# ── DOCX ─────────────────────────────────────────────────────────────────────

def anonymize_docx(
    input_path: str | Path,
    output_path: str | Path,
    mapping: dict[str, str],
    modes: dict[str, str] | None = None,
):
    from docx import Document as DocxDocument

    input_path = Path(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument(input_path)

    # 1. Reemplazo profundo via XPath (Cubre cuerpo, cuadros de texto, dibujos, etc.)
    _deep_replace_xml(doc._element, mapping, modes)

    # 2. Reemplazo profundo en todos los headers/footers
    for section in doc.sections:
        for hf in [section.header, section.footer, section.first_page_header,
                   section.first_page_footer, section.even_page_header, section.even_page_footer]:
            if hf:
                _deep_replace_xml(hf._element, mapping, modes)

    doc.save(str(output_path))


def deanonymize_docx(
    input_path: str | Path,
    output_path: str | Path,
    mapping_path: str | Path,
):
    """
    Restore an anonymized DOCX to its original text using an Excel mapping file.

    The reverse mapping (pseudonym → original) is read from the Excel file
    provided by the user. All entries use substring match mode because 
    pseudonym tokens (e.g. [PERSONA_1]) are unique strings that
    will not cause false partial matches.
    """
    from docx import Document as DocxDocument
    from anonymizer import mapping as map_module

    input_path = Path(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    reverse_mapping = map_module.load_reverse_mapping(mapping_path)
    # All pseudonyms use substring mode — word-boundary fails on bracket tokens
    reverse_modes = {k: "substring" for k in reverse_mapping}

    doc = DocxDocument(input_path)
    _deep_replace_xml(doc._element, reverse_mapping, reverse_modes)

    for section in doc.sections:
        for hf in [section.header, section.footer, section.first_page_header,
                   section.first_page_footer, section.even_page_header, section.even_page_footer]:
            if hf:
                _deep_replace_xml(hf._element, reverse_mapping, reverse_modes)

    doc.save(str(output_path))


def _deep_replace_xml(element, mapping: dict[str, str], modes: dict[str, str] | None = None):
    """
    Busca reemplazos en los párrafos (w:p) para manejar texto fragmentado entre runs (w:r/w:t).
    Luego busca en nodos w:t sueltos para objetos que no sean párrafos estándar.
    """
    from docx.text.paragraph import Paragraph as DocxParagraph
    
    # 1. Procesar párrafos (mejor para texto corrido como 'Juan Pérez')
    for p_node in element.xpath('.//w:p'):
        # Creamos una instancia de Paragraph de python-docx para usar su método .text
        # Pero ojo, no queremos usar el constructor estándar si no tenemos el parent.
        # Una forma más segura es operar directamente sobre los runs del párrafo.
        para = DocxParagraph(p_node, None) 
        full_text = para.text
        if full_text:
            new_text = _apply_mapping_to_text(full_text, mapping, modes)
            if new_text != full_text:
                _set_paragraph_text_safely(para, new_text)

    # 2. Procesar nodos w:t que NO estén dentro de párrafos ya procesados (ej: en tablas complejas o dibujos)
    # Para evitar doble reemplazo, solo si el texto cambió.
    for t_node in element.xpath('.//w:t'):
        if t_node.text:
            new_text = _apply_mapping_to_text(t_node.text, mapping, modes)
            if new_text != t_node.text:
                t_node.text = new_text


def _set_paragraph_text_safely(para, new_text):
    """
    Intenta preservar el formato del primer run si es posible, 
    o simplemente reemplaza el texto total.
    """
    if not para.runs:
        para.add_run(new_text)
        return

    # Estrategia: Ponemos todo el nuevo texto en el primer run y vaciamos los demás
    # para no perder el estilo del párrafo.
    para.runs[0].text = new_text
    for i in range(1, len(para.runs)):
        para.runs[i].text = ""


def _replace_in_paragraph(para, mapping: dict[str, str], modes: dict[str, str] | None = None):
    # Ya no se usa para el flujo DOCX principal porque _deep_replace_xml es más exhaustivo
    pass


# ── PDF ───────────────────────────────────────────────────────────────────────

def anonymize_pdf(
    input_path: str | Path,
    output_path: str | Path,
    mapping: dict[str, str],
    modes: dict[str, str] | None = None,
):
    """
    Extracts text from input PDF, applies mapping, and generates a new PDF
    with the anonymized text using reportlab.
    """
    import pdfplumber
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    input_path = Path(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Extract text page by page
    pages_text: list[list[str]] = []
    with pdfplumber.open(str(input_path)) as pdf:
        for page in pdf.pages:
            raw = page.extract_text() or ""
            lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
            pages_text.append(lines)

    # Build reportlab document
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    normal.fontName = "Helvetica"
    normal.fontSize = 11
    normal.leading = 15

    page_break_style = ParagraphStyle(
        "PageBreak", parent=normal, spaceBefore=20, spaceAfter=20
    )

    story = []
    for page_idx, lines in enumerate(pages_text):
        if page_idx > 0:
            from reportlab.platypus import PageBreak
            story.append(PageBreak())

        for line in lines:
            anonymized_line = _apply_mapping_to_text(line, mapping, modes)
            # Escape XML special chars for reportlab
            safe = anonymized_line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, normal))
            story.append(Spacer(1, 2))

    doc.build(story)


def deanonymize_pdf(
    input_path: str | Path,
    output_path: str | Path,
    mapping_path: str | Path,
):
    """
    Restore an anonymized PDF to its original text using an Excel mapping file.

    The reverse mapping (pseudonym → original) is read from the Excel file. 
    Output is a plain-structural PDF (same limitation as anonymize_pdf).
    """
    import pdfplumber
    from anonymizer import mapping as map_module
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    input_path = Path(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    reverse_mapping = map_module.load_reverse_mapping(mapping_path)
    # All pseudonyms use substring mode (see deanonymize_docx rationale)
    reverse_modes = {k: "substring" for k in reverse_mapping}

    pages_text: list[list[str]] = []
    with pdfplumber.open(str(input_path)) as pdf:
        for page in pdf.pages:
            raw = page.extract_text() or ""
            lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
            pages_text.append(lines)

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    normal.fontName = "Helvetica"
    normal.fontSize = 11
    normal.leading = 15

    ParagraphStyle("PageBreak", parent=normal, spaceBefore=20, spaceAfter=20)

    story = []
    for page_idx, lines in enumerate(pages_text):
        if page_idx > 0:
            from reportlab.platypus import PageBreak
            story.append(PageBreak())

        for line in lines:
            restored_line = _apply_mapping_to_text(line, reverse_mapping, reverse_modes)
            safe = restored_line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, normal))
            story.append(Spacer(1, 2))

    doc.build(story)
