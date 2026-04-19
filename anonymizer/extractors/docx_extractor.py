from pathlib import Path
from docx import Document as DocxDocument
from docx.table import Table
from anonymizer.models import Document


def _iter_block_items(doc_obj):
    """Yield paragraphs and tables in document order (including in cells)."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph

    parent = doc_obj.element.body
    for child in parent.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield DocxParagraph(child, doc_obj)
        elif tag == "tbl":
            yield Table(child, doc_obj)


def _extract_paragraphs(doc_obj):
    """
    Extract all text paragraphs preserving table cell content.
    Returns list of strings, one per paragraph.
    """
    paragraphs = []

    for block in _iter_block_items(doc_obj):
        if hasattr(block, "runs"):
            # It's a paragraph
            text = block.text.strip()
            if text:
                paragraphs.append(text)
        elif hasattr(block, "rows"):
            # It's a table — iterate cells
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            paragraphs.append(text)

    return paragraphs


def _extract_all_xml_text(doc_obj):
    """
    Recorre todo el XML del documento buscando elementos de texto (w:t),
    incluyendo cuadros de texto y formas que no están en el flujo normal de párrafos.
    """
    from docx.oxml.ns import qn
    
    all_text_parts = []
    # Usamos xpath para encontrar todos los w:p y concatenar su texto interno completo
    # Esto une w:hyperlink y w:r fragmentados dentro del mismo párrafo
    for p_node in doc_obj._element.xpath('.//w:p'):
        text = "".join(t.text for t in p_node.xpath('.//w:t') if t.text).strip()
        if text:
            all_text_parts.append(text)
            
    # Buscamos nodos 'w:t' (texto) huérfanos que no estén dentro de un w:p ya procesado
    for t_node in doc_obj._element.xpath('.//w:t'):
        # chequear si pertenece a un w:p
        if not t_node.xpath('ancestor::w:p'):
            if t_node.text and t_node.text.strip():
                all_text_parts.append(t_node.text.strip())
                
    return all_text_parts

def extract(path: str | Path) -> Document:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    path = Path(path)
    doc_obj = DocxDocument(str(path))

    # 1. Extracción tradicional (para mantener orden de párrafos en el proceso de reemplazo si fuera necesario)
    paragraphs = _extract_paragraphs(doc_obj)
    
    # 2. Agregar texto de todos los headers/footers
    for section in doc_obj.sections:
        for hf in [section.header, section.footer, section.first_page_header, 
                   section.first_page_footer, section.even_page_header, section.even_page_footer]:
            if hf:
                for p_node in hf._element.xpath('.//w:p'):
                    text = "".join(t.text for t in p_node.xpath('.//w:t') if t.text).strip()
                    if text:
                        paragraphs.append(text)
                for t_node in hf._element.xpath('.//w:t'):
                    if not t_node.xpath('ancestor::w:p') and t_node.text and t_node.text.strip():
                        paragraphs.append(t_node.text.strip())

    # 3. Búsqueda profunda en el cuerpo (Cuadros de texto, formas, etc.)
    deep_text = _extract_all_xml_text(doc_obj)
    
    # 4. XMLs Auxiliares (Comments, Footnotes, Endnotes)
    aux_text = []
    for rel in doc_obj.part.rels.values():
        if rel.reltype in (RT.COMMENTS, RT.FOOTNOTES, RT.ENDNOTES):
            part = rel.target_part
            if hasattr(part, "element"):
                for p_node in part.element.xpath('.//w:p'):
                    text = "".join(t.text for t in p_node.xpath('.//w:t') if t.text).strip()
                    if text:
                        aux_text.append(text)

    # Unificamos para el full_text que va a la IA (ner)
    seen = set()
    combined = []
    for p in (paragraphs + deep_text + aux_text):
        if p not in seen:
            combined.append(p)
            seen.add(p)

    full_text = "\n".join(combined)

    return Document(path=str(path), paragraphs=combined, full_text=full_text)
